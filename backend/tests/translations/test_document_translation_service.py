# Copyright 2026 Google LLC
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.
"""Tests for the document translation job service (upload/review/export)."""

import datetime
import io
from unittest.mock import AsyncMock, MagicMock

import pytest
from docx import Document
from fastapi import HTTPException

from src.translations.documents import locale_format
from src.translations.documents import service as svc
from src.translations.documents.dto.document_translation_dto import (
    FinalizeUploadDto,
    GenerateUploadUrlDto,
    StartTranslationDto,
    UpdateSegmentDto,
)
from src.translations.documents.service import MAX_UPLOAD_BYTES
from src.translations.documents.schema.document_translation_model import (
    DocumentTranslationJobModel,
    DocumentTranslationSegmentModel,
)
from src.translations.documents.service import DocumentTranslationService


def _fixture_docx_bytes() -> bytes:
    doc = Document()
    doc.add_heading("2.19 Right-of-use assets", level=1)
    doc.add_paragraph("The Group recognised an impairment of EUR 1,234.")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _table_docx_bytes() -> bytes:
    """A heading, a prose line and a financial table with locked figures."""
    doc = Document()
    doc.add_heading("2.19 Right-of-use assets", level=1)
    doc.add_paragraph(
        "Total assets amounted to 319,915 as at January 31, 2026."
    )
    table = doc.add_table(rows=2, cols=2)
    rows = [
        ["€ in thousands", "January 31, 2026"],
        ["Total assets", "319,915"],
    ]
    for r, values in enumerate(rows):
        for c, value in enumerate(values):
            table.rows[r].cells[c].paragraphs[0].add_run(value)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _service() -> DocumentTranslationService:
    jobs = AsyncMock()
    segments = AsyncMock()
    memory = AsyncMock()
    gcs = MagicMock()
    gcs.bucket_name = "test-bucket"
    signer = MagicMock()
    signer.generate_v4_upload_signed_url.return_value = (
        "https://storage.example/signed",
        "gs://test-bucket/document-translations/uploads/u1/report.docx",
    )
    return DocumentTranslationService(
        jobs=jobs, segments=segments, memory=memory, gcs=gcs, signer=signer
    )


def _job_model(**overrides) -> DocumentTranslationJobModel:
    defaults = dict(
        id="job-1",
        filename="report.docx",
        status="review",
        target_market="NL",
        source_gcs_uri="gs://test-bucket/document-translations/job-1/source.docx",
    )
    defaults.update(overrides)
    return DocumentTranslationJobModel(**defaults)


@pytest.mark.anyio
async def test_create_job_rejects_non_docx():
    service = _service()
    with pytest.raises(HTTPException) as exc:
        await service.create_job("report.pdf", b"%PDF-", "n@hp.com")
    assert exc.value.status_code == 400
    assert "Word source" in exc.value.detail


@pytest.mark.anyio
async def test_upload_url_rejects_non_docx_before_minting():
    service = _service()
    with pytest.raises(HTTPException) as exc:
        await service.generate_upload_url(
            GenerateUploadUrlDto(filename="report.pdf", size_bytes=1000)
        )
    assert exc.value.status_code == 400
    service.signer.generate_v4_upload_signed_url.assert_not_called()


@pytest.mark.anyio
async def test_upload_url_rejects_oversized_file():
    service = _service()
    with pytest.raises(HTTPException) as exc:
        await service.generate_upload_url(
            GenerateUploadUrlDto(
                filename="report.docx", size_bytes=MAX_UPLOAD_BYTES + 1
            )
        )
    assert exc.value.status_code == 413


@pytest.mark.anyio
async def test_upload_url_is_minted_for_a_large_docx():
    """The branded FY25-26 report is 55MB — past Cloud Run's body limit."""
    service = _service()
    response = await service.generate_upload_url(
        GenerateUploadUrlDto(
            filename="report.docx", size_bytes=55 * 1024 * 1024
        )
    )
    assert response.upload_url == "https://storage.example/signed"
    assert response.gcs_uri.endswith("report.docx")
    args = service.signer.generate_v4_upload_signed_url.call_args.args
    assert args[1] == (
        "application/vnd.openxmlformats-officedocument"
        ".wordprocessingml.document"
    )
    assert args[2] == "test-bucket"


@pytest.mark.anyio
async def test_finalize_registers_the_uploaded_file_without_reuploading():
    service = _service()
    service.jobs.create.side_effect = lambda model: model
    service.gcs.download_bytes_from_gcs.return_value = _fixture_docx_bytes()
    uri = "gs://test-bucket/document-translations/uploads/u1/report.docx"

    job = await service.finalize_upload(
        FinalizeUploadDto(gcs_uri=uri, filename="report.docx"), "n@hp.com"
    )

    assert job.source_gcs_uri == uri
    assert job.stats["translatable"] == 2
    service.gcs.upload_bytes_to_gcs.assert_not_called()
    service.segments.bulk_create.assert_awaited_once()


@pytest.mark.anyio
async def test_finalize_reports_a_missing_upload():
    service = _service()
    service.gcs.download_bytes_from_gcs.return_value = None
    with pytest.raises(HTTPException) as exc:
        await service.finalize_upload(
            FinalizeUploadDto(gcs_uri="gs://b/gone.docx", filename="a.docx"),
            "n@hp.com",
        )
    assert exc.value.status_code == 400
    assert "upload it again" in exc.value.detail


@pytest.mark.anyio
async def test_create_job_parses_stores_and_persists_segments():
    service = _service()
    service.jobs.create.side_effect = lambda model: model

    job = await service.create_job(
        "report.docx", _fixture_docx_bytes(), "n@hp.com"
    )

    assert job.status == "uploaded"
    assert job.stats["translatable"] == 2  # heading + prose
    chapter = job.stats["chapters"][0]
    assert chapter["title"] == "2.19 Right-of-use assets"
    assert chapter["id"] == "2.19"
    service.gcs.upload_bytes_to_gcs.assert_called_once()
    rows = service.segments.bulk_create.call_args.args[0]
    assert {r.status for r in rows} == {"pending"}
    assert rows[0].job_id == job.id


@pytest.mark.anyio
async def test_retranslate_segment_applies_instruction():
    service = _service()
    service.jobs.get_by_id.return_value = _job_model()
    service.segments.find_by_job.return_value = [
        DocumentTranslationSegmentModel(
            id=1,
            job_id="job-1",
            seg_index=1,
            kind="prose",
            source_text="The Group recognised an impairment.",
            translation="oud",
            status="translated",
        )
    ]
    fake_translator = MagicMock()
    fake_translator.translate_batch.return_value = {1: "nieuw"}
    service._build_translator = AsyncMock(return_value=fake_translator)
    service._load_glossary = AsyncMock(return_value=([], []))

    await service.retranslate_segment("job-1", 1, "more formal")

    service._build_translator.assert_awaited_once()
    assert service._build_translator.await_args.kwargs.get("instruction") == (
        "more formal"
    ) or "more formal" in service._build_translator.await_args.args
    values = service.segments.update_segment.call_args.args[2]
    assert values == {
        "translation": "nieuw",
        "status": "translated",
        "provenance": "ai",
        "finding": None,
    }


@pytest.mark.anyio
async def test_update_segment_rejects_unknown_status():
    service = _service()
    service.jobs.get_by_id.return_value = _job_model()
    with pytest.raises(HTTPException) as exc:
        await service.update_segment(
            "job-1", 3, UpdateSegmentDto(status="wonky")
        )
    assert exc.value.status_code == 400


@pytest.mark.anyio
async def test_update_segment_edit_marks_edited():
    service = _service()
    service.jobs.get_by_id.return_value = _job_model()
    service.segments.update_segment.return_value = (
        DocumentTranslationSegmentModel(
            id=1,
            job_id="job-1",
            seg_index=3,
            kind="prose",
            source_text="x",
            translation="y",
            status="edited",
        )
    )
    await service.update_segment(
        "job-1", 3, UpdateSegmentDto(translation="y")
    )
    values = service.segments.update_segment.call_args.args[2]
    assert values == {
        "translation": "y",
        "provenance": "edited",
        "status": "translated",
    }


@pytest.mark.anyio
async def test_export_blocked_by_qa_errors():
    service = _service()
    service.jobs.get_by_id.return_value = _job_model()
    service.segments.find_by_job.return_value = [
        DocumentTranslationSegmentModel(
            id=1,
            job_id="job-1",
            seg_index=1,
            kind="prose",
            source_text="Total assets were 319,915.",
            translation="De totale activa bedroegen 319.915.",
            status="translated",
            finding={"severity": "error", "type": "number"},
        )
    ]
    with pytest.raises(HTTPException) as exc:
        await service.export("job-1")
    assert exc.value.status_code == 409
    assert "QA findings" in exc.value.detail


@pytest.mark.anyio
async def test_export_blocked_by_an_approved_segment_too():
    """Approving a section does not look at findings, so gating on open
    segments alone would make bulk approval a way around the checks."""
    service = _service()
    service.jobs.get_by_id.return_value = _job_model()
    service.segments.find_by_job.return_value = [
        DocumentTranslationSegmentModel(
            id=1,
            job_id="job-1",
            seg_index=1,
            kind="prose",
            source_text="Total assets were 319,915.",
            translation="De totale activa bedroegen 400.000.",
            status="approved",
            finding={"severity": "error", "type": "number"},
        )
    ]
    with pytest.raises(HTTPException) as exc:
        await service.export("job-1")
    assert exc.value.status_code == 409


@pytest.mark.anyio
async def test_export_is_not_blocked_by_the_report_of_an_earlier_run():
    """`qa_findings` on the job is the report of the run that produced it and
    is never rewritten during review. Gating on it kept refusing an export
    whose findings had all been resolved — a permanent block."""
    service = _service()
    service.jobs.get_by_id.return_value = _job_model(
        qa_findings=[{"severity": "error", "type": "number"}]
    )
    service.gcs.download_bytes_from_gcs.return_value = _fixture_docx_bytes()
    service.segments.find_by_job.return_value = [
        DocumentTranslationSegmentModel(
            id=1,
            job_id="job-1",
            seg_index=1,
            kind="prose",
            source_text="The Group recognised an impairment of EUR 1,234.",
            translation="De Groep verwerkte een waardedaling van EUR 1,234.",
            status="approved",
            finding=None,
        )
    ]

    name, data = await service.export("job-1")

    assert name.endswith("(NL).docx")
    assert data


@pytest.mark.anyio
async def test_export_applies_reviewed_translations():
    service = _service()
    service.jobs.get_by_id.return_value = _job_model()
    service.gcs.download_bytes_from_gcs.return_value = _fixture_docx_bytes()
    service.segments.find_by_job.return_value = [
        DocumentTranslationSegmentModel(
            id=1,
            job_id="job-1",
            seg_index=1,
            kind="prose",
            source_text="The Group recognised an impairment of EUR 1,234.",
            translation=(
                "De Groep heeft een bijzondere waardevermindering van "
                "EUR 1,234 opgenomen."
            ),
            status="approved",
        ),
        DocumentTranslationSegmentModel(
            id=2,
            job_id="job-1",
            seg_index=0,
            kind="heading",
            source_text="2.19 Right-of-use assets",
            translation=None,  # untranslated: must stay English
            status="pending",
        ),
    ]

    filename, data = await service.export("job-1")

    assert filename == "report (NL).docx"
    exported = Document(io.BytesIO(data))
    texts = [p.text for p in exported.paragraphs]
    assert any("bijzondere waardevermindering" in t for t in texts)
    assert "2.19 Right-of-use assets" in texts
    service.gcs.upload_bytes_to_gcs.assert_called_once()
    update_values = service.jobs.update.call_args.args[1]
    assert update_values["status"] == "completed"


# ── resuming an interrupted run ──────────────────────────────────────────
#
# The worker is an asyncio task inside one Cloud Run instance and the service
# scales to zero, so an instance reclaimed mid-run takes the task with it
# WITHOUT raising: the job keeps status "translating" at its last flushed
# percentage. A real 1,300-segment annual report froze at 36% that way, and the
# 409 on "translating" meant it could never be started again either.


def _stale(minutes: int) -> datetime.datetime:
    return datetime.datetime.now(datetime.UTC) - datetime.timedelta(
        minutes=minutes
    )


def _seg(index: int, section: str, translation: str | None, status: str):
    return DocumentTranslationSegmentModel(
        id=index,
        job_id="job-1",
        seg_index=index,
        kind="prose",
        section_id=section,
        section_path=[section],
        source_text=f"source {index}",
        translation=translation,
        status=status,
    )


class TestStallDetection:
    def test_a_beating_job_is_not_stalled(self):
        job = _job_model(status="translating", updated_at=_stale(minutes=1))

        assert svc._is_stalled(job) is False

    def test_a_silent_translating_job_is_stalled(self):
        job = _job_model(status="translating", updated_at=_stale(minutes=45))

        assert svc._is_stalled(job) is True

    def test_only_translating_jobs_can_stall(self):
        """A finished job is old by definition — that is not a stall."""
        for state in ("review", "completed", "failed", "uploaded"):
            job = _job_model(status=state, updated_at=_stale(minutes=600))
            assert svc._is_stalled(job) is False, state


class TestResumeGuards:
    @pytest.mark.anyio
    async def test_a_live_run_is_not_resumable(self):
        service = _service()
        service.jobs.get_by_id.return_value = _job_model(
            status="translating", updated_at=_stale(minutes=1)
        )

        with pytest.raises(HTTPException) as exc:
            await service.resume_translation("job-1")

        assert exc.value.status_code == 409
        assert "still translating" in exc.value.detail

    @pytest.mark.anyio
    async def test_a_finished_run_has_nothing_to_resume(self):
        service = _service()
        service.jobs.get_by_id.return_value = _job_model(status="review")

        with pytest.raises(HTTPException) as exc:
            await service.resume_translation("job-1")

        assert exc.value.status_code == 409

    @pytest.mark.anyio
    async def test_starting_over_is_allowed_once_a_run_has_stalled(self):
        """The 409 that wedged the stuck job must not fire on a dead one."""
        service = _service()
        service.jobs.get_by_id.return_value = _job_model(
            status="translating", updated_at=_stale(minutes=45)
        )
        service.segments.find_by_job.return_value = []
        service.memory.find_matches.return_value = {}
        service.jobs.update.return_value = _job_model(status="translating")
        service._build_translator = AsyncMock()
        service._launch = MagicMock()

        await service.start_translation(
            "job-1", StartTranslationDto(target_market="NL")
        )

        assert service._launch.called

    @pytest.mark.anyio
    async def test_resume_keeps_the_market_and_model_of_the_dead_run(self):
        service = _service()
        service.jobs.get_by_id.return_value = _job_model(
            status="translating",
            target_market="DE",
            model_id="gemini-x",
            updated_at=_stale(minutes=45),
        )
        service.jobs.update.return_value = _job_model(status="translating")
        service._build_translator = AsyncMock()
        service._launch = MagicMock()

        await service.resume_translation("job-1")

        service._build_translator.assert_awaited_once_with("DE", "gemini-x")
        assert service._launch.call_args.kwargs["resume"] is True
        # The progress dict is untouched, so the bar carries on from where it
        # froze instead of snapping back to zero.
        assert "progress" not in service.jobs.update.call_args.args[1]

    @pytest.mark.anyio
    async def test_resume_keeps_the_notation_choice_of_the_dead_run(self):
        """Notation is decided when a run starts. Resuming with a different
        answer would renotate half a document and leave the rest English."""
        service = _service()
        service.jobs.get_by_id.return_value = _job_model(
            status="translating",
            target_market="DE",
            localise_numbers=True,
            updated_at=_stale(minutes=45),
        )
        service.jobs.update.return_value = _job_model(status="translating")
        service._build_translator = AsyncMock()
        service._launch = MagicMock()

        await service.resume_translation("job-1")

        assert (
            service._launch.call_args.args[2]
            == locale_format.for_market("DE")
        )

    @pytest.mark.anyio
    async def test_a_plain_run_resumes_without_renotation(self):
        service = _service()
        service.jobs.get_by_id.return_value = _job_model(
            status="translating",
            target_market="DE",
            localise_numbers=False,
            updated_at=_stale(minutes=45),
        )
        service.jobs.update.return_value = _job_model(status="translating")
        service._build_translator = AsyncMock()
        service._launch = MagicMock()

        await service.resume_translation("job-1")

        assert service._launch.call_args.args[2] is None

    @pytest.mark.anyio
    async def test_starting_a_run_hands_the_notation_to_the_worker(self):
        """QA has to accept the localised spelling of a figure, so the worker
        needs the same format the export will use."""
        service = _service()
        service.jobs.get_by_id.return_value = _job_model(status="uploaded")
        service.segments.find_by_job.return_value = []
        service.memory.find_matches.return_value = {}
        service.jobs.update.return_value = _job_model(status="translating")
        service._build_translator = AsyncMock()
        service._launch = MagicMock()

        await service.start_translation(
            "job-1",
            StartTranslationDto(target_market="NL", localise_numbers=True),
        )

        assert (
            service._launch.call_args.args[2]
            == locale_format.for_market("NL")
        )
        assert (
            service.jobs.update.call_args.args[1]["localise_numbers"] is True
        )

    @pytest.mark.anyio
    async def test_resume_needs_a_market_to_resume_with(self):
        service = _service()
        service.jobs.get_by_id.return_value = _job_model(
            status="failed", target_market=None
        )

        with pytest.raises(HTTPException) as exc:
            await service.resume_translation("job-1")

        assert exc.value.status_code == 400


class _FakeSession:
    """Stands in for `async_session_local()` — the worker opens its own."""

    async def __aenter__(self):
        return "db"

    async def __aexit__(self, *_exc):
        return False


class TestResumeRun:
    """What resume actually has to get right, at the worker level."""

    def _wire(self, monkeypatch, rows_before, rows_after):
        seg_repo, job_repo = AsyncMock(), AsyncMock()
        seg_repo.find_by_job.side_effect = [rows_before, rows_after]
        monkeypatch.setattr(svc, "async_session_local", lambda: _FakeSession())
        monkeypatch.setattr(
            svc, "DocumentTranslationSegmentRepository", lambda _db: seg_repo
        )
        monkeypatch.setattr(
            svc, "DocumentTranslationJobRepository", lambda _db: job_repo
        )
        qa_run = MagicMock(return_value=[])
        monkeypatch.setattr(svc.qa, "run_all", qa_run)
        translator = MagicMock()
        translator.translate_batch.side_effect = lambda batch: {
            s.id: f"vertaling {s.id}" for s in batch
        }
        return seg_repo, job_repo, qa_run, translator

    @pytest.mark.anyio
    async def test_only_the_untranslated_segments_go_back_to_the_model(
        self, monkeypatch
    ):
        # Section "1" was finished by the dead run; section "2" never got there.
        before = [
            _seg(1, "1", "vertaling 1", "translated"),
            _seg(2, "1", "vertaling 2", "translated"),
            _seg(3, "2", None, "pending"),
            _seg(4, "2", None, "pending"),
        ]
        after = [
            _seg(i, s, f"vertaling {i}", "translated")
            for i, s in ((1, "1"), (2, "1"), (3, "2"), (4, "2"))
        ]
        seg_repo, job_repo, qa_run, translator = self._wire(
            monkeypatch, before, after
        )

        await svc.DocumentTranslationService(
            jobs=AsyncMock(), segments=AsyncMock(), memory=AsyncMock(),
            gcs=MagicMock(), signer=MagicMock(),
        )._run_translation("job-1", translator, resume=True)

        sent = [
            s.id
            for call in translator.translate_batch.call_args_list
            for s in call.args[0]
        ]
        assert sent == [3, 4], "already-translated segments were re-sent"

    @pytest.mark.anyio
    async def test_progress_carries_on_instead_of_restarting_at_zero(
        self, monkeypatch
    ):
        before = [
            _seg(1, "1", "vertaling 1", "translated"),
            _seg(2, "1", "vertaling 2", "translated"),
            _seg(3, "2", None, "pending"),
        ]
        after = [_seg(i, s, f"vertaling {i}", "translated")
                 for i, s in ((1, "1"), (2, "1"), (3, "2"))]
        seg_repo, job_repo, qa_run, translator = self._wire(
            monkeypatch, before, after
        )

        await svc.DocumentTranslationService(
            jobs=AsyncMock(), segments=AsyncMock(), memory=AsyncMock(),
            gcs=MagicMock(), signer=MagicMock(),
        )._run_translation("job-1", translator, resume=True)

        progress = [
            c.args[1]["progress"]
            for c in job_repo.update.call_args_list
            if "progress" in c.args[1]
        ]
        assert progress, "no progress was flushed"
        assert progress[0]["total"] == 3
        # 2 carried + the 1 just done — not 1 out of 1.
        assert progress[-1]["translated"] == 3
        # And the finished section is not shown as queued again.
        assert progress[-1]["sections"]["1"] == "done"

    @pytest.mark.anyio
    async def test_qa_judges_the_whole_document_not_just_this_pass(
        self, monkeypatch
    ):
        """`set_findings` clears every prior finding, so a resume that ran QA
        over its own pass only would erase the earlier half's findings."""
        before = [
            _seg(1, "1", "vertaling 1", "translated"),
            _seg(2, "2", None, "pending"),
        ]
        after = [
            _seg(1, "1", "vertaling 1", "translated"),
            _seg(2, "2", "vertaling 2", "translated"),
        ]
        seg_repo, job_repo, qa_run, translator = self._wire(
            monkeypatch, before, after
        )

        await svc.DocumentTranslationService(
            jobs=AsyncMock(), segments=AsyncMock(), memory=AsyncMock(),
            gcs=MagicMock(), signer=MagicMock(),
        )._run_translation("job-1", translator, resume=True)

        judged = sorted(s.id for s in qa_run.call_args.args[0])
        assert judged == [1, 2]

    @pytest.mark.anyio
    async def test_a_restart_still_replaces_everything(self, monkeypatch):
        """Starting over is the other button: prior output is not kept."""
        before = [
            _seg(1, "1", "oude vertaling", "translated"),
            _seg(2, "1", None, "pending"),
        ]
        after = [_seg(1, "1", "vertaling 1", "translated"),
                 _seg(2, "1", "vertaling 2", "translated")]
        seg_repo, job_repo, qa_run, translator = self._wire(
            monkeypatch, before, after
        )

        await svc.DocumentTranslationService(
            jobs=AsyncMock(), segments=AsyncMock(), memory=AsyncMock(),
            gcs=MagicMock(), signer=MagicMock(),
        )._run_translation("job-1", translator, resume=False)

        sent = [
            s.id
            for call in translator.translate_batch.call_args_list
            for s in call.args[0]
        ]
        assert sent == [1, 2]
def _localisation_rows() -> list[DocumentTranslationSegmentModel]:
    """Translations that reproduced the source's figures, as instructed."""
    return [
        DocumentTranslationSegmentModel(
            id=1,
            job_id="job-1",
            seg_index=0,
            kind="heading",
            source_text="2.19 Right-of-use assets",
            translation="2.19 Gebruiksrechten",
            status="approved",
        ),
        DocumentTranslationSegmentModel(
            id=2,
            job_id="job-1",
            seg_index=1,
            kind="prose",
            source_text=(
                "Total assets amounted to 319,915 as at January 31, 2026."
            ),
            translation=(
                "De totale activa bedroegen 319,915 per January 31, 2026."
            ),
            status="approved",
        ),
        DocumentTranslationSegmentModel(
            id=3,
            job_id="job-1",
            seg_index=2,
            kind="table_label",
            source_text="€ in thousands",
            translation="€ in duizenden",
            status="approved",
        ),
        DocumentTranslationSegmentModel(
            id=4,
            job_id="job-1",
            seg_index=3,
            kind="table_label",
            source_text="January 31, 2026",
            translation="January 31, 2026",
            status="approved",
        ),
        DocumentTranslationSegmentModel(
            id=5,
            job_id="job-1",
            seg_index=4,
            kind="table_label",
            source_text="Total assets",
            translation="Totale activa",
            status="approved",
        ),
    ]


async def _export_table_docx(**job_overrides) -> list[str]:
    service = _service()
    service.jobs.get_by_id.return_value = _job_model(**job_overrides)
    service.gcs.download_bytes_from_gcs.return_value = _table_docx_bytes()
    service.segments.find_by_job.return_value = _localisation_rows()

    _, data = await service.export("job-1")

    exported = Document(io.BytesIO(data))
    cells = [
        cell.text
        for table in exported.tables
        for row in table.rows
        for cell in row.cells
    ]
    return [p.text for p in exported.paragraphs] + cells


@pytest.mark.anyio
async def test_export_renotates_figures_and_dates_when_asked():
    texts = await _export_table_docx(localise_numbers=True)

    assert "De totale activa bedroegen 319.915 per 31 januari 2026." in texts
    # The locked figure cell is never translated, only renotated.
    assert "319.915" in texts
    assert "31 januari 2026" in texts
    # A section number is a reference, not a figure.
    assert "2.19 Gebruiksrechten" in texts
    assert not any("319,915" in t for t in texts)


@pytest.mark.anyio
async def test_export_keeps_english_notation_by_default():
    texts = await _export_table_docx()

    assert "De totale activa bedroegen 319,915 per January 31, 2026." in texts
    assert "319,915" in texts
    assert not any("319.915" in t for t in texts)


class TestResolvingAFinding:
    """A blocking finding has to be resolvable without re-running the whole
    document: the reviewer corrects the figure and the flag goes."""

    def _service_with(self, translation: str):
        service = _service()
        service.jobs.get_by_id.return_value = _job_model()
        service.segments.find_by_job.return_value = [
            DocumentTranslationSegmentModel(
                id=1,
                job_id="job-1",
                seg_index=1,
                kind="prose",
                source_text="Total assets were 319,915.",
                translation=translation,
                status="translated",
                finding={"severity": "error", "type": "number"},
            )
        ]
        service.segments.update_segment.return_value = (
            DocumentTranslationSegmentModel(
                id=1,
                job_id="job-1",
                seg_index=1,
                kind="prose",
                source_text="Total assets were 319,915.",
                translation=translation,
                status="translated",
            )
        )
        service._load_glossary = AsyncMock(return_value=([], []))
        return service

    @pytest.mark.anyio
    async def test_correcting_the_figure_clears_the_flag(self):
        service = self._service_with("De totale activa bedroegen 319.915.")

        await service.update_segment(
            "job-1",
            1,
            UpdateSegmentDto(
                translation="De totale activa bedroegen 319,915."
            ),
        )

        values = service.segments.update_segment.call_args.args[2]
        assert values["finding"] is None

    @pytest.mark.anyio
    async def test_an_edit_that_is_still_wrong_keeps_the_flag(self):
        service = self._service_with("De totale activa bedroegen 319.915.")

        await service.update_segment(
            "job-1",
            1,
            UpdateSegmentDto(
                translation="De totale activa bedroegen 400,000."
            ),
        )

        values = service.segments.update_segment.call_args.args[2]
        assert values["finding"]["severity"] == "error"
        assert values["finding"]["type"] == "number"

    @pytest.mark.anyio
    async def test_a_localised_run_accepts_the_reviewers_notation(self):
        """With renotation on, 319.915 is the same figure as 319,915 — the
        export writes it that way anyway."""
        service = self._service_with("De totale activa bedroegen 319.915.")
        service.jobs.get_by_id.return_value = _job_model(
            localise_numbers=True
        )

        await service.update_segment(
            "job-1",
            1,
            UpdateSegmentDto(
                translation="De totale activa bedroegen 319.915."
            ),
        )

        values = service.segments.update_segment.call_args.args[2]
        assert values["finding"] is None

    @pytest.mark.anyio
    async def test_a_retranslation_that_drops_a_figure_is_flagged(self):
        """Clearing the finding blindly handed this a free pass."""
        service = self._service_with("De totale activa waren hoog.")
        fake = MagicMock()
        fake.translate_batch.return_value = {1: "De totale activa waren hoog."}
        service._build_translator = AsyncMock(return_value=fake)

        await service.retranslate_segment("job-1", 1, None)

        values = service.segments.update_segment.call_args.args[2]
        assert values["finding"]["type"] == "number"
        assert values["finding"]["expected"] == "319,915"
