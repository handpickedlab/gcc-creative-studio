# Copyright 2026 Google LLC
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.
"""Tests for the docx document-translation engine (parse/apply/QA)."""

import io

import pytest
from docx import Document

from src.translations.documents import qa
from src.translations.documents.docx_engine import DocxTranslationEngine
from src.translations.documents.model import SegmentKind, classify_text
from src.translations.documents.translator import (
    GlossaryEntry,
    PseudoTranslator,
    iter_batches,
    translate_tree,
)


def _fixture_docx() -> io.BytesIO:
    """A miniature annual report: headings, styled prose, a financial table."""
    doc = Document()
    doc.add_heading("1. Management Board Report", level=1)
    p = doc.add_paragraph("We operate ")
    p.add_run("700 stores").bold = True
    p.add_run(" across 11 countries.")
    doc.add_heading("2.19 Right-of-use assets", level=2)
    doc.add_paragraph(
        "The Group recognised an impairment of EUR 1,234 thousand."
    )
    table = doc.add_table(rows=3, cols=3)
    rows = [
        ["€ in thousands", "January 31, 2026", "January 31, 2025"],
        ["Intangible assets", "20,913", "22,210"],
        ["Total assets", "319,915", "-"],
    ]
    for r, values in enumerate(rows):
        for c, value in enumerate(values):
            table.rows[r].cells[c].paragraphs[0].add_run(value)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _merged_table_docx() -> io.BytesIO:
    """A table whose first row is merged across all three columns."""
    doc = Document()
    doc.add_heading("2.22 Inventories", level=1)
    table = doc.add_table(rows=3, cols=3)
    for r, values in enumerate(
        [
            ["Inventories", "", ""],
            ["Goods for resale", "63,900", "59,847"],
            ["Total inventories", "63,900", "59,847"],
        ]
    ):
        for c, value in enumerate(values):
            if value:
                table.rows[r].cells[c].paragraphs[0].add_run(value)
    table.rows[0].cells[0].merge(table.rows[0].cells[2])
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _text_box_docx() -> io.BytesIO:
    """A paragraph with its own text that also hosts a text box.

    Real documents wrap the box in VML/DrawingML; the parser only looks for
    ``w:txbxContent`` at any depth, so the fixture keeps the wrapper minimal.
    """
    from docx.oxml.ns import qn
    from docx.oxml.parser import OxmlElement

    doc = Document()
    doc.add_heading("1.3 Environmental and Social", level=1)
    host = doc.add_paragraph("Our targets are verified externally.")

    holder = OxmlElement("w:r")
    box = OxmlElement("w:txbxContent")
    inner_p = OxmlElement("w:p")
    inner_r = OxmlElement("w:r")
    inner_t = OxmlElement("w:t")
    inner_t.text = "Scope 1 emissions fell by 12%."
    inner_r.append(inner_t)
    inner_p.append(inner_r)
    box.append(inner_p)
    holder.append(box)
    host._p.append(holder)
    assert host._p.find(f".//{qn('w:txbxContent')}") is not None

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _tracked_changes_docx() -> io.BytesIO:
    from docx.oxml.ns import nsmap
    from docx.oxml.parser import parse_xml

    doc = Document()
    p = doc.add_paragraph("Deferred tax assets ")
    ns = f'xmlns:w="{nsmap["w"]}"'
    p._p.append(
        parse_xml(
            f'<w:ins {ns} w:id="1" w:author="a"><w:r><w:t>have been offset.'
            f"</w:t></w:r></w:ins>"
        )
    )
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


class TestClassify:
    @pytest.mark.parametrize(
        "text,expected",
        [
            ("319,915", SegmentKind.NUMERIC),
            ("(141,764)", SegmentKind.NUMERIC),
            ("19%", SegmentKind.NUMERIC),
            ("2.18", SegmentKind.NUMERIC),
            ("-", SegmentKind.SKIP),
            ("", SegmentKind.SKIP),
            ("Intangible assets", SegmentKind.TABLE_LABEL),
        ],
    )
    def test_table_cells(self, text, expected):
        assert classify_text(text, in_table=True, is_heading=False) == expected

    def test_prose_and_heading(self):
        assert (
            classify_text("We operate.", in_table=False, is_heading=False)
            == SegmentKind.PROSE
        )
        assert (
            classify_text("2.19 Right-of-use", in_table=False, is_heading=True)
            == SegmentKind.HEADING
        )

    def test_date_with_digits_is_translatable(self):
        assert (
            classify_text("January 31, 2026", in_table=True, is_heading=False)
            == SegmentKind.TABLE_LABEL
        )


class TestParse:
    def test_tree_structure_and_classification(self):
        engine = DocxTranslationEngine(_fixture_docx())
        tree = engine.tree

        top = tree.root.children
        assert [s.title for s in top] == ["1. Management Board Report"]
        assert [c.title for c in top[0].children] == [
            "2.19 Right-of-use assets"
        ]

        stats = tree.stats()
        # 2 headings + 2 prose + 5 table labels (incl. the date columns)
        # are translatable; 3 numeric cells and the "-" cell are not.
        assert stats["heading"] == 2
        assert stats["prose"] == 2
        assert stats["table_label"] == 5
        assert stats["numeric"] == 3
        assert stats["translatable"] == 9

    def test_table_and_row_indexes(self):
        engine = DocxTranslationEngine(_fixture_docx())
        cells = [s for s in engine.tree.segments if s.table_index is not None]
        assert {s.table_index for s in cells} == {0}
        assert [s.row_index for s in cells] == [0, 0, 0, 1, 1, 1, 2, 2, 2]
        prose = [s for s in engine.tree.segments if s.kind == SegmentKind.PROSE]
        assert all(s.table_index is None for s in prose)

    def test_section_path_inherited(self):
        engine = DocxTranslationEngine(_fixture_docx())
        prose = [
            s for s in engine.tree.segments if s.kind == SegmentKind.PROSE
        ]
        assert prose[1].section_path == (
            "1. Management Board Report",
            "2.19 Right-of-use assets",
        )


class TestApplyRoundTrip:
    def test_translation_lands_and_numbers_survive(self):
        engine = DocxTranslationEngine(_fixture_docx())
        failed = translate_tree(engine.tree.translatable, PseudoTranslator())
        assert failed == []
        applied = engine.apply()
        assert applied == 9

        out = io.BytesIO()
        engine.save(out)
        out.seek(0)
        reopened = DocxTranslationEngine(out)

        texts = [s.text for s in reopened.tree.segments]
        assert "«We operate 700 stores across 11 countries.»" in texts
        assert "«Intangible assets»" in texts
        # Numeric cells were never touched.
        assert "319,915" in texts
        assert "20,913" in texts
        # Same shape: no elements created or destroyed.
        assert len(reopened.tree.segments) == len(engine.tree.segments)
        assert len(reopened.document.tables) == 1

    def test_multi_run_paragraph_keeps_run_count(self):
        engine = DocxTranslationEngine(_fixture_docx())
        seg = next(
            s for s in engine.tree.segments if "700 stores" in s.text
        )
        seg.translation = "Wij exploiteren 700 winkels in 11 landen."
        engine.apply()
        # 3 runs still exist; text is consolidated into the first.
        runs = seg.paragraph.runs
        assert len(runs) == 3
        assert runs[0].text == "Wij exploiteren 700 winkels in 11 landen."
        assert runs[1].text == "" and runs[2].text == ""


class TestMergedCells:
    def test_merged_cell_counted_once_and_parse_is_deterministic(self):
        """Regression: keying the seen-set on id() let recycled ids collide,
        so unrelated cells were silently dropped from the parse."""
        runs = []
        for _ in range(5):
            engine = DocxTranslationEngine(_merged_table_docx())
            stats = engine.tree.stats()
            runs.append((stats["total"], stats["translatable"]))
        assert len(set(runs)) == 1, f"parse is not deterministic: {set(runs)}"

        engine = DocxTranslationEngine(_merged_table_docx())
        labels = [
            s.text
            for s in engine.tree.segments
            if s.kind == SegmentKind.TABLE_LABEL
        ]
        assert labels.count("Inventories") == 1
        assert labels.count("Total inventories") == 1

    def test_every_cell_of_a_plain_table_is_parsed(self):
        engine = DocxTranslationEngine(_fixture_docx())
        cells = [s for s in engine.tree.segments if s.table_index is not None]
        assert len(cells) == 9  # 3x3, none merged, none empty


class TestTextBoxes:
    def test_text_box_content_becomes_its_own_segment(self):
        engine = DocxTranslationEngine(_text_box_docx())
        texts = [s.text for s in engine.tree.segments]
        assert "Our targets are verified externally." in texts
        assert "Scope 1 emissions fell by 12%." in texts

    def test_host_translation_does_not_clobber_the_text_box(self):
        """Regression: apply() searched text nodes recursively, so a host
        paragraph's translation overwrote the text box nested inside it."""
        engine = DocxTranslationEngine(_text_box_docx())
        translate_tree(engine.tree.translatable, PseudoTranslator())
        engine.apply()
        out = io.BytesIO()
        engine.save(out)
        out.seek(0)

        reopened = DocxTranslationEngine(out)
        texts = [s.text for s in reopened.tree.segments]
        assert "«Our targets are verified externally.»" in texts
        assert "«Scope 1 emissions fell by 12%.»" in texts


class TestTrackedChanges:
    def test_tracked_changes_are_counted_for_preflight(self):
        engine = DocxTranslationEngine(_tracked_changes_docx())
        assert engine.tracked_changes() == 1

    def test_clean_document_reports_none(self):
        engine = DocxTranslationEngine(_fixture_docx())
        assert engine.tracked_changes() == 0


class TestOutline:
    def test_explicit_numbers_in_the_headings_win(self):
        engine = DocxTranslationEngine(_fixture_docx())
        chapters = engine.tree.outline()
        assert [c["id"] for c in chapters] == ["1"]
        assert [s["id"] for s in chapters[0]["sections"]] == ["2.19"]
        assert chapters[0]["sections"][0]["tables"] == 1

    def test_position_supplies_the_number_when_the_text_lacks_one(self):
        """Word usually renders section numbers from automatic numbering, so
        the heading text itself carries none."""
        doc = Document()
        doc.add_heading("Management Board Report", level=1)
        doc.add_paragraph("Body.")
        doc.add_heading("CEO Statement", level=2)
        doc.add_paragraph("More body.")
        doc.add_heading("Brand", level=2)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        chapters = DocxTranslationEngine(buffer).tree.outline()
        assert [c["id"] for c in chapters] == ["1"]
        assert [s["id"] for s in chapters[0]["sections"]] == ["1.1", "1.2"]

    def test_segments_carry_their_review_section(self):
        engine = DocxTranslationEngine(_fixture_docx())
        prose = [
            s for s in engine.tree.segments if s.kind == SegmentKind.PROSE
        ]
        assert prose[0].section_id == "1"
        assert prose[1].section_id == "2.19"
        cells = [s for s in engine.tree.segments if s.table_index is not None]
        assert {s.section_id for s in cells} == {"2.19"}

    def test_heading_levels_and_bold_are_captured(self):
        engine = DocxTranslationEngine(_fixture_docx())
        headings = [
            s for s in engine.tree.segments if s.kind == SegmentKind.HEADING
        ]
        assert [s.heading_level for s in headings] == [1, 2]
        engine2 = DocxTranslationEngine(_merged_table_docx())
        assert not any(
            s.bold for s in engine2.tree.segments if s.text == "Inventories"
        )


class TestBatching:
    def test_batches_respect_section_boundaries(self):
        engine = DocxTranslationEngine(_fixture_docx())
        batches = list(iter_batches(engine.tree.translatable))
        for batch in batches:
            assert len({s.section_path for s in batch}) == 1

    def test_batches_respect_size_bounds(self):
        engine = DocxTranslationEngine(_fixture_docx())
        batches = list(
            iter_batches(engine.tree.translatable, max_segments=2)
        )
        assert all(len(b) <= 2 for b in batches)


class TestQa:
    def _translated_fixture(self):
        engine = DocxTranslationEngine(_fixture_docx())
        translate_tree(engine.tree.translatable, PseudoTranslator())
        return engine.tree.translatable

    def test_clean_translation_passes(self):
        segments = self._translated_fixture()
        assert qa.run_all(segments) == []

    def test_dropped_number_is_an_error(self):
        segments = self._translated_fixture()
        seg = next(s for s in segments if "1,234" in s.text)
        seg.translation = "De Groep heeft een waardevermindering geboekt."
        findings = qa.check_numbers(segments)
        assert [f.check for f in findings] == ["number"]
        assert findings[0].severity == qa.Severity.ERROR
        assert "1,234" in findings[0].detail
        assert findings[0].expected == "1,234"

    def test_invented_number_is_an_error(self):
        segments = self._translated_fixture()
        seg = next(s for s in segments if "700 stores" in s.text)
        seg.translation = "«We operate 800 stores across 11 countries.»"
        findings = qa.check_numbers(segments)
        assert any("800" in f.detail for f in findings)

    def test_translated_brand_name_is_an_error(self):
        segments = self._translated_fixture()
        seg = next(s for s in segments if "impairment" in s.text)
        seg.translation = seg.translation.replace("EUR", "euro")
        findings = qa.check_do_not_translate(segments, ["EUR"])
        assert findings and findings[0].check == "dnt"
        assert findings[0].term == "EUR"

    def test_glossary_miss_is_a_warning(self):
        segments = self._translated_fixture()
        glossary = [
            GlossaryEntry(
                source="impairment", target="bijzondere waardevermindering"
            )
        ]
        findings = qa.check_glossary(segments, glossary)
        assert findings and findings[0].severity == qa.Severity.WARNING
        assert findings[0].term == "impairment"
        assert findings[0].expected == "bijzondere waardevermindering"

    def test_length_blowup_is_a_warning(self):
        segments = self._translated_fixture()
        seg = next(s for s in segments if "1,234" in s.text)
        seg.translation = seg.text + " lorem ipsum" * 20
        findings = qa.check_length(segments)
        assert findings and findings[0].check == "length"
